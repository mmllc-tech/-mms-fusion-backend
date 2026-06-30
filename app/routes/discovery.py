"""
Discovery — file upload, text extraction, AI extraction via Claude API
Supports: XLSX, DOCX, PDF, TXT, CSV
"""
import os, json, io, re
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from app.models.database import get_db, Project, UploadedFile
import httpx

router = APIRouter()

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL = "claude-sonnet-4-6"

REQUIRED_QUESTIONS = [
    "legislation","enterprise","legal_entity","ldg",
    "payroll_freq","salary_basis","elements",
    "worker_types","positions","modules_enabled","security"
]

EXTRACTION_PROMPT = """You are an expert Oracle HCM Core HR implementation consultant.

Read these discovery documents and extract answers to the following Oracle HCM implementation questions.
Return ONLY a valid JSON object with the question IDs as keys.
Only include questions where you found clear evidence — do not guess or invent.

Questions to answer:
- "legislation": Country and legislation (e.g. United States, US Legislation, USD)
- "enterprise": Enterprise / company name in Oracle HCM
- "legal_entity": Legal Entities and Legal Employers in scope, with EIN numbers if found
- "ldg": Legislative Data Group name (e.g. US Legislative Data Group)
- "bus_units": Business Units in scope
- "payroll_freq": Payroll frequency (Biweekly, Monthly, Semimonthly etc)
- "salary_basis": Salary bases required (Annual, Hourly, annualization factors)
- "elements": Earning elements in scope for conversion
- "payroll_int": Payroll interface or integration requirements
- "worker_types": Worker types in scope (Employee, Contingent Worker etc)
- "positions": Whether Position Management is enabled and its scope
- "modules_enabled": Oracle modules enabled for this implementation
- "security": Custom security roles required with their access scope
- "dff": Descriptive Flexfield segments needed beyond standard

Documents:
{content}

Return only JSON: {{"legislation": "...", "enterprise": "...", ...}}
Only include keys where you found clear evidence."""

def extract_text_from_excel(content: bytes, filename: str) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
        text = f"Excel file: {filename}\nSheets: {', '.join(wb.sheetnames)}\n\n"
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_text = f"--- Sheet: {sheet_name} ---\n"
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(v) if v is not None else "" for v in row]
                if any(v.strip() for v in row_vals):
                    sheet_text += " | ".join(row_vals) + "\n"
            text += sheet_text + "\n"
        return text[:50000]
    except Exception as e:
        return f"Excel file: {filename} (parse error: {str(e)})"

def extract_text_from_docx(content: bytes, filename: str) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        text = f"Document: {filename}\n\n"
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
        return text[:50000]
    except Exception as e:
        return f"Document: {filename} (parse error: {str(e)})"

def extract_text_from_pdf(content: bytes, filename: str) -> str:
    try:
        import pdfplumber
        text = f"PDF: {filename}\n\n"
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text[:50000]
    except Exception as e:
        return f"PDF: {filename} (parse error: {str(e)})"

def extract_text(content: bytes, filename: str) -> str:
    ext = filename.split(".")[-1].lower()
    if ext in ("xlsx", "xls"):
        return extract_text_from_excel(content, filename)
    elif ext == "docx":
        return extract_text_from_docx(content, filename)
    elif ext == "pdf":
        return extract_text_from_pdf(content, filename)
    else:
        try:
            return content.decode("utf-8", errors="replace")[:50000]
        except:
            return f"File: {filename}"

async def call_claude(content: str, api_key: str) -> dict:
    prompt =
